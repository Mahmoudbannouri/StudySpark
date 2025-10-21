import os
import sys
from typing import List

try:
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
except ImportError:
    print("Missing dependency: python-docx. Install with: pip install python-docx", file=sys.stderr)
    sys.exit(1)


def read_text(path: str) -> List[str]:
    with open(path, 'r', encoding='utf-8') as f:
        return f.read().splitlines()


def add_code_block(doc: Document, lines: List[str]):
    # Add code block lines as a single preformatted paragraph with monospace font
    para = doc.add_paragraph()
    run = para.add_run('\n'.join(lines))
    run.font.name = 'Consolas'
    # Ensure proper font mapping for docx
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
    run.font.size = Pt(10)


def add_heading(doc: Document, text: str, level: int):
    level = max(1, min(6, level))
    doc.add_heading(text.strip(), level=level)


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(text.strip())
    p.style = 'List Bullet'


def add_paragraph(doc: Document, text: str):
    doc.add_paragraph(text)


def export_markdown_to_docx(md_path: str, docx_path: str):
    lines = read_text(md_path)
    doc = Document()
    doc.core_properties.title = os.path.basename(md_path)

    in_code = False
    code_lines: List[str] = []

    for line in lines:
        # Code block handling
        if line.strip().startswith('```'):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                # closing code fence
                add_code_block(doc, code_lines)
                in_code = False
                code_lines = []
            continue

        if in_code:
            code_lines.append(line)
            continue

        # Headings (# to ######)
        if line.startswith('#'):
            hashes = len(line) - len(line.lstrip('#'))
            text = line.lstrip('#').strip()
            add_heading(doc, text, hashes)
            continue

        # Bulleted list
        if line.lstrip().startswith(('- ', '* ')):
            # strip leading spaces and marker
            t = line.lstrip()[2:]
            add_bullet(doc, t)
            continue

        # Horizontal rule or empty line -> blank paragraph
        if line.strip() == '' or set(line.strip()) == set('-'):
            add_paragraph(doc, '')
            continue

        # Default paragraph
        add_paragraph(doc, line)

    # If file ended inside a code block
    if in_code and code_lines:
        add_code_block(doc, code_lines)

    # Save
    os.makedirs(os.path.dirname(docx_path) or '.', exist_ok=True)
    doc.save(docx_path)


def main():
    # Resolve default paths relative to repo root
    repo_root = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..'))
    md_path = os.path.join(repo_root, 'StudySpark_AI_Conversation_and_Code_Log.md')
    out_path = os.path.join(repo_root, 'StudySpark_AI_Conversation_and_Code_Log.docx')

    if len(sys.argv) >= 2:
        md_path = sys.argv[1]
        if not os.path.isabs(md_path):
            md_path = os.path.abspath(md_path)
    if len(sys.argv) >= 3:
        out_path = sys.argv[2]
        if not os.path.isabs(out_path):
            out_path = os.path.abspath(out_path)

    if not os.path.exists(md_path):
        print(f"Markdown file not found: {md_path}", file=sys.stderr)
        sys.exit(2)

    export_markdown_to_docx(md_path, out_path)
    print(f"Exported to: {out_path}")


if __name__ == '__main__':
    main()
